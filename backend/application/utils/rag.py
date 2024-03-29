from typing import List, Union
from pathlib import Path
from PyPDF2 import PdfReader
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain_community.vectorstores import Chroma
from langchain.embeddings import OpenAIEmbeddings
from langchain import hub
from langchain_community.chat_models import ChatOpenAI
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
# from langchain_core.runnables import RunnableParallel
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.messages import HumanMessage
from langchain_core.output_parsers import StrOutputParser
from langchain.vectorstores import Qdrant
# from langchain.memory import ConversationBufferMemory
# from langchain.prompts import PromptTemplate
from typing import IO, Union, Optional
from pathlib import Path
from backend.application.utils.utils import get_hash_value
from docx import Document as Docxreader

import requests
from bs4 import BeautifulSoup
from hashlib import md5

import os

class DocReader:

    def __init__(self, file: Optional[Union[IO[bytes], Path]] = None):
        self.file_path = file
        self.name = file.name if file else None
        self.chat_history = []
        self.llm = ChatOpenAI(model_name="gpt-3.5-turbo", temperature=0)

    def pdf_loader(self) -> List[Document]:
        # read the pdf file
        pdf_reader = PdfReader(self.file_path)
        pages = []
        # loop through the pages with page number
        for idx, page in enumerate(pdf_reader.pages):
            text = page.extract_text()
            pages.append(text)
        # join all the pages into one string
        pages = "\n\n".join(pages)
        # calculate hash value of the content
        hash_value = get_hash_value(pages)
        # create a document object
        doc = [Document(page_content=pages, metadata={"source": str(self.name), "hash": hash_value})]
        return doc
    
    def docx_loader(self) -> List[Document]:
        # read the Word document
        docx_reader = Docxreader(self.file_path)
        paragraphs = []
        # loop through the paragraphs
        for paragraph in docx_reader.paragraphs:
            text = paragraph.text
            paragraphs.append(text)
        # join all the paragraphs into one string
        paragraphs = "\n\n".join(paragraphs)
        # calculate hash value of the content
        hash_value = get_hash_value(paragraphs)
        # create a document object
        doc = [Document(page_content=paragraphs, metadata={"source": str(self.name), "hash": hash_value})]
        return doc
    
    def txt_loader(self) -> List[Document]:
        # read the text file
        self.file_path.seek(0)  # ensure you're at the start of the file
        content = self.file_path.read().decode()  # decode from bytes to string if necessary
        # calculate hash value of the content
        hash_value = get_hash_value(content)
        # create a document object
        doc = [Document(page_content=content, metadata={"source": str(self.name), "hash": hash_value})]
        return doc
    
    def url_loader(self, url: str) -> List[Document]:
        response = requests.get(url)
        if response.status_code == 200 and 'text/html' in response.headers['Content-Type']:
            soup = BeautifulSoup(response.content, 'html.parser')
            text = soup.get_text(separator='\n')
            hash_value = md5(text.encode()).hexdigest()
            return [Document(page_content=text, metadata={"source": url, "hash": hash_value})]
        else:
            return []

    def split_text(self, doc: Document, chunk_size: int = 1000, chunk_overlap: int = 200):
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap, add_start_index=True)
        print("My doc is: ", type(doc))
        splits = text_splitter.split_documents(doc)
        return splits
    
    def create_vector_store(self,splits: List[Document]):
        # vectorstore = Chroma.from_documents(documents=splits, embedding=OpenAIEmbeddings(disallowed_special=()))
        vectorstore = Qdrant.from_documents(
                            splits,
                            OpenAIEmbeddings(disallowed_special=()),
                            url=os.getenv("QDRANT_URL"),
                            prefer_grpc=True,
                            api_key=os.getenv("QDRANT_API_KEY"),
                            collection_name="my_documents",
                        )
        return vectorstore

    def create_retriever(self, vectorstore: Qdrant):
        retriever = vectorstore.as_retriever(search_type="similarity_score_threshold", search_kwargs={"score_threshold": 0.5})
        return retriever
    
    def pull_prompt(self):
        prompt = hub.pull("rlm/rag-prompt")
        return prompt
    
    def create_llm(self):
        # llm = ChatOpenAI(model_name="gpt-3.5-turbo", temperature=0)
        return self.llm
    
    def format_docs(self, docs: List[Document]):
        return "\n\n".join(doc.page_content for doc in docs)
    
    ############################ Contextual Functions ############################

    def create_qa_prompt(self):
        qa_system_prompt = """You are an assistant for question-answering tasks. \
        Use the following pieces of retrieved context to answer the question. \
        If you don't know the answer, just say that you don't know. \
        Use three sentences maximum and keep the answer concise.\

        {context}"""

        
        qa_prompt = ChatPromptTemplate.from_messages(
            [
                ("system", qa_system_prompt),
                *self.chat_history,
                ("human", "{question}"),
            ]
        )
        return qa_prompt


    def contextualized_question(self, input: dict):
        contextualize_q_system_prompt = """Given a chat history and the latest user question \
        which might reference context in the chat history, formulate a standalone question \
        which can be understood without the chat history. Do NOT answer the question, \
        just reformulate it if needed and otherwise return it as is."""

        contextualize_q_prompt = ChatPromptTemplate.from_messages(
            [
                ("system", contextualize_q_system_prompt),
                MessagesPlaceholder(variable_name="chat_history"),
                ("human", "{question}"),
            ]
        )
        contextualize_q_chain = contextualize_q_prompt | self.llm | StrOutputParser()

        if input.get("chat_history"):
            return contextualize_q_chain.invoke(input)
        else:
            return input["question"]

############################ Contextual Functions ############################

    ### Create RAG chain
    
    def create_rag_chain(self, retriever, prompt, llm):
        rag_chain = (
            {"context": retriever | self.format_docs, "question": RunnablePassthrough()}
            | prompt
            | llm
            | StrOutputParser()
        )
        return rag_chain
    
    def create_rag_chain_with_history(self,contextualized_question, retriever, format_docs, qa_prompt, llm):
        rag_chain = (
            RunnablePassthrough.assign(
                context=contextualized_question | retriever | format_docs
            )
            | qa_prompt
            | llm
        )
        return rag_chain
    
    # ask question
    
    def ask_question(self, rag_chain, question: str):
        response = rag_chain.invoke(question)
        return response
    
    def ask_contextual_question(self, rag_chain, question: str):
        """
        Asks a question to the RAG model and updates the chat history.

        Args:
        question (str): The question to be asked.

        Returns:
        str: The answer from the RAG model.
        """
        ai_msg = rag_chain.invoke({"question": question, "chat_history": self.chat_history})
        self.chat_history.extend([HumanMessage(content=question), ai_msg])
        return ai_msg.content

if __name__ == "__main__":
    file_path = Path("yolo.pdf")
    doc_reader = DocReader(file_path)
    doc = doc_reader.pdf_loader()
    # splits = doc_reader.split_text(doc)
    # vectorstore = doc_reader.create_vector_store(splits)
    # retriever = doc_reader.create_retriever(vectorstore)
    # llm = doc_reader.create_llm()
    # prompt = doc_reader.create_qa_prompt()
    # rag_chain = doc_reader.create_rag_chain_with_history(doc_reader.contextualized_question, retriever, doc_reader.format_docs, prompt, llm)
    # response_1 = doc_reader.ask_contextual_question(rag_chain, "What is Yolov7?")
    # response_2 = doc_reader.ask_contextual_question(rag_chain, "What are it's benefits?")
    # print(response_1)
    # print(response_2)
    print(doc)
   
    
    
    
    